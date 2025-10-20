"""
Markdown to Word Document Converter
====================================

This script converts the Markdown analysis report to a professional
Microsoft Word document (.docx) format.

Features:
- Parses Markdown headings, lists, tables, and images
- Professional styling with custom fonts and colors
- Embedded images from reports/assets/
- Japanese text support
- Table of contents
- Headers and footers with page numbers

Author: PDF Reporting Specialist
Date: October 2025
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background_color(cell, color):
    """
    Set background color for a table cell.

    Args:
        cell: Table cell object
        color: RGB color tuple (r, g, b)
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '%02x%02x%02x' % color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.

    Args:
        paragraph: Paragraph object
        url: URL string
        text: Display text
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def setup_styles(doc):
    """
    Configure custom styles for the document.

    Args:
        doc: Document object
    """
    # Heading 1 style
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Calibri'
    heading1.font.size = Pt(24)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(10, 64, 115)  # Dark blue
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(12)

    # Heading 2 style
    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Calibri'
    heading2.font.size = Pt(18)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(10, 64, 115)
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(8)
    heading2.paragraph_format.keep_with_next = True

    # Heading 3 style
    heading3 = doc.styles['Heading 3']
    heading3.font.name = 'Calibri'
    heading3.font.size = Pt(14)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(44, 95, 141)
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)

    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(8)


def parse_markdown_table(table_text):
    """
    Parse Markdown table into rows and columns.

    Args:
        table_text: Markdown table as string

    Returns:
        List of lists (rows and columns)
    """
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]

    # Remove separator line (second line with dashes)
    if len(lines) > 1 and re.match(r'^[\|\-\s:]+$', lines[1]):
        lines.pop(1)

    # Parse rows
    rows = []
    for line in lines:
        # Remove leading/trailing pipes and split by pipe
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        rows.append(cells)

    return rows


def add_table_to_doc(doc, rows):
    """
    Add a formatted table to the document.

    Args:
        doc: Document object
        rows: List of lists (table data)
    """
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    # Fill in data first
    for i, row_data in enumerate(rows):
        cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            cells[j].text = cell_data

            # Format text if cell has content
            if cells[j].paragraphs[0].runs:
                cells[j].paragraphs[0].runs[0].font.size = Pt(10)

            # Align numbers to the right
            if re.match(r'^[\d,\.¥\-%]+$', cell_data.strip()):
                cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Format header row (after text is added)
    header_cells = table.rows[0].cells
    for cell in header_cells:
        set_cell_background_color(cell, (52, 152, 219))  # Blue background
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()  # Add spacing after table


def convert_markdown_to_docx(md_path, output_path, assets_dir):
    """
    Convert Markdown file to Word document.

    Args:
        md_path: Path to Markdown file
        output_path: Path for output .docx file
        assets_dir: Directory containing image assets
    """
    # Read Markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    doc = Document()

    # Set up page
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4 height
    section.page_width = Cm(21)     # A4 width
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Setup styles
    setup_styles(doc)

    # Add header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Multi-Store Fashion Retail Sales Analysis - January 2024"
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Add footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_para.text = "Page "

    # Add page number field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Parse Markdown content
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_buffer = []

    while i < len(lines):
        line = lines[i]

        # Skip empty lines unless in table
        if not line.strip() and not in_table:
            i += 1
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('##'):
            doc.add_heading(line[2:].strip(), level=1)

        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)

        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)

        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)

        # Horizontal rule
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 80)
            p.runs[0].font.color.rgb = RGBColor(200, 200, 200)

        # Images
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text, img_path = match.groups()
                full_img_path = Path(assets_dir) / img_path.replace('assets/', '')

                if full_img_path.exists():
                    doc.add_paragraph(alt_text, style='Heading 4')
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    run.add_picture(str(full_img_path), width=Inches(5.5))
                    doc.add_paragraph()  # Spacing after image

        # Tables (start)
        elif line.startswith('|') and not in_table:
            in_table = True
            table_buffer = [line]

        # Tables (continuation)
        elif in_table:
            if line.startswith('|'):
                table_buffer.append(line)
            else:
                # End of table
                table_text = '\n'.join(table_buffer)
                rows = parse_markdown_table(table_text)
                add_table_to_doc(doc, rows)
                in_table = False
                table_buffer = []
                continue  # Don't increment i, process current line again

        # Bullet lists
        elif re.match(r'^\d+\.\s', line):  # Numbered list
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.5)

        elif line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)

        # Bold text
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Normal paragraph
        else:
            doc.add_paragraph(line.strip())

        i += 1

    # Handle remaining table if document ends with one
    if in_table and table_buffer:
        table_text = '\n'.join(table_buffer)
        rows = parse_markdown_table(table_text)
        add_table_to_doc(doc, rows)

    # Save document
    doc.save(output_path)
    print(f"✅ Word document created: {output_path}")
    print(f"   File size: {Path(output_path).stat().st_size / 1024:.2f} KB")


def main():
    """
    Main execution function.
    """
    # Define paths
    project_root = Path(__file__).parent.parent.parent
    md_file = project_root / 'reports' / 'analysis_report.md'
    output_file = project_root / 'reports' / 'detailed_report.docx'
    assets_dir = project_root / 'reports'

    print("=" * 80)
    print("MARKDOWN TO WORD CONVERSION")
    print("=" * 80)
    print(f"Input file:  {md_file}")
    print(f"Output file: {output_file}")
    print(f"Assets dir:  {assets_dir}")
    print()

    # Check input file exists
    if not md_file.exists():
        print(f"❌ Error: Markdown file not found: {md_file}")
        return False

    # Convert
    print("Converting Markdown to Word document...")
    convert_markdown_to_docx(md_file, output_file, assets_dir)

    print()
    print("=" * 80)
    print("✅ CONVERSION COMPLETE")
    print("=" * 80)
    print(f"\nWord document ready: {output_file}")
    print("\nYou can now:")
    print("  - Open in Microsoft Word")
    print("  - Edit and add comments")
    print("  - Share with stakeholders")
    print("  - Print for distribution")

    return True


if __name__ == "__main__":
    import sys
    success = main()
    sys.exit(0 if success else 1)
